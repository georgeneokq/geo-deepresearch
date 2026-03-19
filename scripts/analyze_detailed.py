#!/usr/bin/env python3
"""
DOCX Detailed Analyzer - Inspects paragraph formatting in detail.
"""

import sys
from pathlib import Path
from docx import Document


def analyze_paragraph_detailed(paragraph, idx):
    """Analyze a paragraph's formatting in detail."""
    info = {
        'index': idx,
        'text': paragraph.text[:80] + '...' if len(paragraph.text) > 80 else paragraph.text,
        'style': paragraph.style.name if paragraph.style else 'None',
        'alignment': paragraph.alignment.name if paragraph.alignment else 'None',
        'is_heading': False,
        'heading_level': None,
        'has_numbering': False,
        'font_name': None,
        'font_size': None,
        'bold': False,
        'jc_alignment': None,  # Raw XML alignment
    }
    
    # Check raw XML for jc (justification) element
    pPr = paragraph._element.pPr
    if pPr is not None:
        jc = pPr.jc
        if jc is not None:
            info['jc_alignment'] = jc.val
    
    # Check if it's a heading style
    if paragraph.style.name.startswith('Heading'):
        info['is_heading'] = True
        info['heading_level'] = int(paragraph.style.name.split()[-1])
    
    # Check for numbering
    if pPr is not None:
        numPr = pPr.numPr
        if numPr is not None:
            info['has_numbering'] = True
    
    # Get font info from first run
    if paragraph.runs:
        run = paragraph.runs[0]
        info['font_name'] = run.font.name
        info['font_size'] = run.font.size.pt if run.font.size else None
        info['bold'] = run.font.bold
    
    return info


def analyze_document_detailed(doc_path: str, start_idx: int = 0):
    """Analyze a Word document starting from a specific paragraph."""
    doc = Document(doc_path)
    
    print("=" * 80)
    print(f"DETAILED ANALYSIS: {doc_path}")
    print(f"Starting from paragraph: {start_idx}")
    print("=" * 80)
    
    headings_without_numbering = []
    centered_headings = []
    
    for idx in range(start_idx, len(doc.paragraphs)):
        paragraph = doc.paragraphs[idx]
        info = analyze_paragraph_detailed(paragraph, idx)
        
        if info['is_heading']:
            if not info['has_numbering']:
                headings_without_numbering.append(info)
            if info['alignment'] == 'CENTER':
                centered_headings.append(info)
            
            indent = "  " * (info['heading_level'] - 1) if info['heading_level'] else ""
            numbering_status = "✓" if info['has_numbering'] else "✗"
            print(f"[{idx}] {indent}H{info['heading_level']}: {info['text'][:60]}")
            print(f"      Numbering: {numbering_status}, Alignment: {info['alignment']}, JC: {info['jc_alignment']}, Font: {info['font_name']}, Size: {info['font_size']}pt")
    
    print("\n" + "=" * 80)
    print("ISSUES FOUND")
    print("=" * 80)
    
    print(f"\nHeadings without numbering: {len(headings_without_numbering)}")
    for info in headings_without_numbering[:20]:
        print(f"  [{info['index']}] H{info['heading_level']}: {info['text'][:50]}...")
    
    print(f"\nCentered headings: {len(centered_headings)}")
    for info in centered_headings[:20]:
        print(f"  [{info['index']}] H{info['heading_level']}: {info['text'][:50]}...")
    
    return headings_without_numbering, centered_headings


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python analyze_detailed.py <document.docx> [start_idx]")
        sys.exit(1)
    
    doc_path = sys.argv[1]
    start_idx = int(sys.argv[2]) if len(sys.argv) > 2 else 0
    
    analyze_document_detailed(doc_path, start_idx)
