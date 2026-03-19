#!/usr/bin/env python3
"""
DOCX Formatter - Standardizes font, size, color, and numbering in a Word document.

This script:
1. Sets all body text to Arial 11pt black
2. Sets headings to Arial with proper sizes (H1: 16pt, H2: 14pt, H3: 13pt, H4: 12pt)
3. Ensures all headings have proper multilevel numbering
4. Fixes heading hierarchy (converts misclassified long H1s to body text)
"""

import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# Configuration
FONT_NAME = "Arial"
BODY_FONT_SIZE = 11
FONT_COLOR = RGBColor(0, 0, 0)  # Black

# All headings standardized to 13pt
HEADING_SIZES = {
    1: 13,
    2: 13,
    3: 13,
    4: 13,
}

HEADING_SPACE_BEFORE = 12
HEADING_SPACE_AFTER = 6
BODY_SPACE_BEFORE = 6
BODY_SPACE_AFTER = 6

# Threshold for considering a heading as misclassified (too long)
HEADING_MAX_LENGTH = 80

# Marker text that indicates the start of main content
MAIN_CONTENT_MARKER = "Capstone Project Final Report for the Project"


def set_run_font(run, font_name=FONT_NAME, font_size=BODY_FONT_SIZE, 
                 color=FONT_COLOR, bold=False, italic=False):
    """Apply font styling to a run."""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.color.rgb = color
    run.font.bold = bold
    run.font.italic = italic
    # Set East Asian font as well
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def set_paragraph_numbering(paragraph, num_id, level):
    """Add numbering to a paragraph."""
    pPr = paragraph._element.get_or_add_pPr()
    
    # Remove existing numbering if present
    existing_numPr = pPr.find(qn('w:numPr'))
    if existing_numPr is not None:
        pPr.remove(existing_numPr)
    
    # Create numPr element
    numPr = OxmlElement('w:numPr')
    
    # Create ilvl (indentation level)
    ilvl = OxmlElement('w:ilvl')
    ilvl.set(qn('w:val'), str(level))
    numPr.append(ilvl)
    
    # Create numId (numbering ID)
    numId = OxmlElement('w:numId')
    numId.set(qn('w:val'), str(num_id))
    numPr.append(numId)
    
    pPr.append(numPr)


def create_multilevel_numbering(doc):
    """Create a multilevel numbering definition for headings."""
    numbering_part = doc.part.numbering_part
    numbering_xml = numbering_part._element
    
    # Find existing abstractNum IDs to avoid conflicts
    existing_abstract_nums = numbering_xml.findall(
        './/w:abstractNum', 
        namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    )
    abstract_num_id = len(existing_abstract_nums) + 1
    
    # Create the abstractNum element
    abstract_num = OxmlElement('w:abstractNum')
    abstract_num.set(qn('w:abstractNumId'), str(abstract_num_id))
    
    # MultiLevelType
    multi_level_type = OxmlElement('w:multiLevelType')
    multi_level_type.set(qn('w:val'), 'multilevel')
    abstract_num.append(multi_level_type)
    
    # Define levels 0-3 (for Heading 1-4)
    level_formats = [
        ('%1.', 'decimal', 0),
        ('%1.%2.', 'decimal', 1),
        ('%1.%2.%3.', 'decimal', 2),
        ('%1.%2.%3.%4.', 'decimal', 3),
    ]
    
    for level_idx, (fmt, num_fmt, restart) in enumerate(level_formats):
        lvl = OxmlElement('w:lvl')
        lvl.set(qn('w:ilvl'), str(level_idx))
        lvl.set(qn('w:tplc'), '04090001')
        
        # Number format
        num_fmt_elem = OxmlElement('w:numFmt')
        num_fmt_elem.set(qn('w:val'), num_fmt)
        lvl.append(num_fmt_elem)
        
        # Level restart
        lvl_restart = OxmlElement('w:lvlRestart')
        lvl_restart.set(qn('w:val'), '1')
        lvl.append(lvl_restart)
        
        # Level text
        lvl_text = OxmlElement('w:lvlText')
        lvl_text.set(qn('w:val'), fmt)
        lvl.append(lvl_text)
        
        # Paragraph properties for this level
        pPr = OxmlElement('w:pPr')
        ind = OxmlElement('w:ind')
        ind.set(qn('w:left'), str(720 * level_idx))
        ind.set(qn('w:hanging'), '360')
        pPr.append(ind)
        lvl.append(pPr)
        
        # Run properties for this level
        rPr = OxmlElement('w:rPr')
        fonts = OxmlElement('w:rFonts')
        fonts.set(qn('w:ascii'), FONT_NAME)
        fonts.set(qn('w:hAnsi'), FONT_NAME)
        fonts.set(qn('w:eastAsia'), FONT_NAME)
        rPr.append(fonts)
        lvl.append(rPr)
        
        abstract_num.append(lvl)
    
    numbering_xml.append(abstract_num)
    
    # Find existing num IDs
    existing_nums = numbering_xml.findall(
        './/w:num', 
        namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    )
    num_id = len(existing_nums) + 1
    
    # Create concrete numbering instance
    num = OxmlElement('w:num')
    num.set(qn('w:numId'), str(num_id))
    
    abstract_num_id_elem = OxmlElement('w:abstractNumId')
    abstract_num_id_elem.set(qn('w:val'), str(abstract_num_id))
    num.append(abstract_num_id_elem)
    
    numbering_xml.append(num)
    numbering_part._numbering = numbering_xml
    
    return num_id


def is_misclassified_heading(paragraph):
    """Check if a heading is likely misclassified (too long, looks like body text)."""
    text = paragraph.text.strip()
    
    # Empty headings are misclassified
    if not text:
        return True
    
    # Very long headings are likely body text
    if len(text) > HEADING_MAX_LENGTH:
        return True
    
    # Headings that look like sentences (end with period and are long)
    if text.endswith('.') and len(text) > 50:
        return True
    
    return False


def fix_heading_style(paragraph, level, num_id):
    """Apply proper heading style to a paragraph."""
    # Apply font to all runs
    for run in paragraph.runs:
        set_run_font(
            run, 
            font_size=HEADING_SIZES.get(level, 13),
            bold=True
        )
    
    # If no runs exist, create one
    if not paragraph.runs:
        run = paragraph.add_run(paragraph.text)
        set_run_font(
            run, 
            font_size=HEADING_SIZES.get(level, 13),
            bold=True
        )
    
    # Set paragraph spacing and alignment
    paragraph.paragraph_format.space_before = Pt(HEADING_SPACE_BEFORE)
    paragraph.paragraph_format.space_after = Pt(HEADING_SPACE_AFTER)
    paragraph.paragraph_format.keep_with_next = level < 3
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Explicitly left-align
    
    # Add numbering
    set_paragraph_numbering(paragraph, num_id, level - 1)


def fix_body_paragraph(paragraph):
    """Apply proper body text style to a paragraph."""
    # Apply font to all runs
    for run in paragraph.runs:
        set_run_font(run, font_size=BODY_FONT_SIZE)
    
    # Set paragraph spacing
    paragraph.paragraph_format.space_before = Pt(BODY_SPACE_BEFORE)
    paragraph.paragraph_format.space_after = Pt(BODY_SPACE_AFTER)


def fix_table(table):
    """Apply proper styling to a table."""
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                if row_idx == 0:  # Header row
                    for run in paragraph.runs:
                        set_run_font(run, font_size=BODY_FONT_SIZE, bold=True)
                else:
                    for run in paragraph.runs:
                        set_run_font(run, font_size=BODY_FONT_SIZE)


def format_document(doc_path: str, output_path: str = None):
    """Format a Word document with consistent styling."""
    if output_path is None:
        output_path = doc_path.replace('.docx', '_formatted.docx')
    
    doc = Document(doc_path)
    
    print(f"Formatting document: {doc_path}")
    print(f"Output: {output_path}")
    print(f"Total paragraphs: {len(doc.paragraphs)}")
    
    # First, find the marker paragraph that indicates start of main content
    main_content_start_idx = None
    for idx, paragraph in enumerate(doc.paragraphs):
        if MAIN_CONTENT_MARKER in paragraph.text:
            main_content_start_idx = idx
            print(f"Found main content marker at paragraph {idx}: {paragraph.text[:60]}...")
            break
    
    if main_content_start_idx is None:
        print("Warning: Main content marker not found. Formatting entire document.")
        main_content_start_idx = 0
    else:
        print(f"Cover pages: paragraphs 0-{main_content_start_idx - 1} (will be preserved)")
        print(f"Main content starts at paragraph {main_content_start_idx}")
    
    # Create multilevel numbering (only applied to main content)
    num_id = create_multilevel_numbering(doc)
    print(f"Created numbering definition with ID: {num_id}")
    
    # Track statistics
    stats = {
        'headings_fixed': 0,
        'body_text_fixed': 0,
        'misclassified_fixed': 0,
        'tables_fixed': 0,
        'cover_pages_preserved': main_content_start_idx if main_content_start_idx > 0 else 0,
    }
    
    # First pass: identify and fix misclassified headings (only in main content)
    print("\n📝 First pass: Fixing misclassified headings...")
    for idx in range(main_content_start_idx, len(doc.paragraphs)):
        paragraph = doc.paragraphs[idx]
        style_name = paragraph.style.name if paragraph.style else ''
        
        # Check if it's a heading style
        if style_name.startswith('Heading'):
            try:
                level = int(style_name.split()[-1])
            except (ValueError, IndexError):
                level = 1
            
            # Check if misclassified
            if is_misclassified_heading(paragraph):
                # Convert to body text by changing style
                paragraph.style = 'Normal'
                stats['misclassified_fixed'] += 1
                print(f"  [{idx}] Converted misclassified H{level} to body text: {paragraph.text[:50]}...")
    
    # Second pass: apply consistent formatting (only in main content)
    print("\n📝 Second pass: Applying consistent formatting...")
    for idx in range(main_content_start_idx, len(doc.paragraphs)):
        paragraph = doc.paragraphs[idx]
        style_name = paragraph.style.name if paragraph.style else ''
        
        if style_name.startswith('Heading'):
            try:
                level = int(style_name.split()[-1])
                level = min(level, 4)  # Cap at H4
            except (ValueError, IndexError):
                level = 1
            
            fix_heading_style(paragraph, level, num_id)
            stats['headings_fixed'] += 1
            
        elif style_name == 'Normal' or not style_name:
            fix_body_paragraph(paragraph)
            stats['body_text_fixed'] += 1
        
        # Handle other paragraph styles as body text
        else:
            fix_body_paragraph(paragraph)
            stats['body_text_fixed'] += 1
    
    # Fix tables (only those in main content)
    print("\n📝 Third pass: Fixing tables...")
    # Note: python-docx doesn't give us table positions easily, so we fix all tables
    # This is acceptable since cover page tables typically already have good styling
    for table in doc.tables:
        fix_table(table)
        stats['tables_fixed'] += 1
    
    # Save the document
    doc.save(output_path)
    
    print("\n" + "=" * 60)
    print("✅ FORMATTING COMPLETE")
    print("=" * 60)
    print(f"Cover pages preserved: {stats['cover_pages_preserved']} paragraphs")
    print(f"Headings fixed: {stats['headings_fixed']}")
    print(f"Body text paragraphs fixed: {stats['body_text_fixed']}")
    print(f"Misclassified headings converted: {stats['misclassified_fixed']}")
    print(f"Tables fixed: {stats['tables_fixed']}")
    print(f"\nSaved to: {output_path}")
    
    return stats


def main():
    if len(sys.argv) < 2:
        print("Usage: python format_docx.py <input.docx> [output.docx]")
        print("\nExample:")
        print("  uv run --with python-docx scripts/format_docx.py docs/final_report.docx")
        sys.exit(1)
    
    input_path = sys.argv[1]
    output_path = sys.argv[2] if len(sys.argv) > 2 else None
    
    try:
        format_document(input_path, output_path)
    except Exception as e:
        print(f"Error: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == "__main__":
    main()
