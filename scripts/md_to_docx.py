#!/usr/bin/env python3
"""
Markdown to DOCX Converter

Converts a Markdown file to a Microsoft Word document (.docx) using python-docx.
Designed for converting architecture documentation with proper styling.

Usage:
    uv run --with python-docx --with markdown --with beautifulsoup4 scripts/md_to_docx.py input.md output.docx

Dependencies:
    - python-docx: For creating Word documents
    - markdown: For parsing Markdown
    - beautifulsoup4: For parsing HTML
"""

# =============================================================================
# CONFIGURATION - Customize these values to change document styling
# =============================================================================

# Font settings
FONT_NAME = "Calibri"  # Default Word font
FONT_SIZE = 11  # Points
FONT_COLOR = "#000000"  # Black

# Heading styles (font size in points)
HEADING_1_SIZE = 16  # Matches sample.docx
HEADING_2_SIZE = 14
HEADING_3_SIZE = 13
HEADING_4_SIZE = 12

# Heading colors
HEADING_1_COLOR = "#000000"  # Black
HEADING_2_COLOR = "#000000"  # Black
HEADING_3_COLOR = "#000000"  # Black
HEADING_4_COLOR = "#000000"  # Black

# Heading bold
HEADING_1_BOLD = True
HEADING_2_BOLD = True
HEADING_3_BOLD = True
HEADING_4_BOLD = True

# Spacing (in points)
PARAGRAPH_SPACE_BEFORE = 6
PARAGRAPH_SPACE_AFTER = 6
HEADING_SPACE_BEFORE = 12
HEADING_SPACE_AFTER = 6

# Code block styling
CODE_FONT_NAME = "Consolas"
CODE_FONT_SIZE = 10
CODE_BACKGROUND_COLOR = "#F5F5F5"  # Light gray

# Table styling
TABLE_HEADER_BG_COLOR = "#4472C4"  # Blue
TABLE_HEADER_FONT_COLOR = "#FFFFFF"  # White
TABLE_BORDER_COLOR = "#000000"  # Black
TABLE_CELL_PADDING = 6  # Points

# Page setup
PAGE_WIDTH = 8.5  # Inches (Letter size)
PAGE_HEIGHT = 11  # Inches
PAGE_MARGIN_TOP = 1  # Inches
PAGE_MARGIN_BOTTOM = 1  # Inches
PAGE_MARGIN_LEFT = 1  # Inches
PAGE_MARGIN_RIGHT = 1  # Inches

# =============================================================================
# END OF CONFIGURATION
# =============================================================================

import sys
import re
from pathlib import Path
from typing import Optional, Tuple

import markdown
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def parse_markdown(md_content: str) -> str:
    """
    Parse Markdown content to HTML.
    """
    md = markdown.Markdown(extensions=[
        'tables',
        'fenced_code',
        'codehilite',
    ])
    return md.convert(md_content)


def hex_to_rgb(hex_color: str) -> RGBColor:
    """Convert hex color code to RGBColor."""
    hex_color = hex_color.lstrip('#')
    r = int(hex_color[0:2], 16)
    g = int(hex_color[2:4], 16)
    b = int(hex_color[4:6], 16)
    return RGBColor(r, g, b)


def set_paragraph_font(paragraph, font_name: str = FONT_NAME, 
                       font_size: int = FONT_SIZE, 
                       font_color: str = FONT_COLOR,
                       bold: bool = False,
                       italic: bool = False):
    """Apply font styling to a paragraph."""
    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.font.color.rgb = hex_to_rgb(font_color)
        run.font.bold = bold
        run.font.italic = italic
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def add_numbering_to_paragraph(paragraph, num_id: int, level: int = 0):
    """
    Add numbering to a paragraph using the specified numbering definition.
    
    Args:
        paragraph: docx paragraph object
        num_id: The numbering definition ID to use
        level: Numbering level (0 for top-level, 1 for sub-level, etc.)
    """
    pPr = paragraph._element.get_or_add_pPr()
    
    # Remove existing numbering if present
    existing_numPr = pPr.find(qn('w:numPr'))
    if existing_numPr is not None:
        pPr.remove(existing_numPr)
    
    # Create numPr element
    numPr = OxmlElement('w:numPr')
    
    # Create ilvl (indentation level)
    ilvl = OxmlElement('w:ilvl')
    ilvl.set(qn('w:val'), str(level))
    numPr.append(ilvl)
    
    # Create numId (numbering ID)
    numId = OxmlElement('w:numId')
    numId.set(qn('w:val'), str(num_id))
    numPr.append(numId)
    
    pPr.append(numPr)


def create_multilevel_numbering(doc: Document) -> int:
    """
    Create a multilevel numbering definition for headings.
    Returns the numId of the created numbering.
    
    Format:
    1. Heading 1
    1.1. Heading 2
    1.1.1. Heading 3
    1.1.1.1. Heading 4
    """
    # Get the numbering element (the actual XML element)
    numbering_part = doc.part.numbering_part
    numbering_xml = numbering_part._element
    
    # Find existing abstractNum IDs to avoid conflicts
    existing_abstract_nums = numbering_xml.findall('.//w:abstractNum', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
    abstract_num_id = len(existing_abstract_nums) + 1
    
    # Create the abstractNum element
    abstract_num = OxmlElement('w:abstractNum')
    abstract_num.set(qn('w:abstractNumId'), str(abstract_num_id))
    
    # MultiLevelType
    multi_level_type = OxmlElement('w:multiLevelType')
    multi_level_type.set(qn('w:val'), 'multilevel')
    abstract_num.append(multi_level_type)
    
    # Define levels 0-3 (for Heading 1-4)
    level_formats = [
        ('%1.', 'decimal', 0),      # Level 0: 1, 2, 3
        ('%1.%2.', 'decimal', 1),   # Level 1: 1.1, 1.2, 1.3
        ('%1.%2.%3.', 'decimal', 2), # Level 2: 1.1.1, 1.1.2
        ('%1.%2.%3.%4.', 'decimal', 3), # Level 3: 1.1.1.1
    ]
    
    for level_idx, (fmt, num_fmt, restart) in enumerate(level_formats):
        lvl = OxmlElement('w:lvl')
        lvl.set(qn('w:ilvl'), str(level_idx))
        
        # Start at 1
        lvl.set(qn('w:tplc'), '04090001')
        
        # Number format
        num_fmt_elem = OxmlElement('w:numFmt')
        num_fmt_elem.set(qn('w:val'), num_fmt)
        lvl.append(num_fmt_elem)
        
        # Level restart
        lvl_restart = OxmlElement('w:lvlRestart')
        lvl_restart.set(qn('w:val'), '1')
        lvl.append(lvl_restart)
        
        # Level text
        lvl_text = OxmlElement('w:lvlText')
        lvl_text.set(qn('w:val'), fmt)
        lvl.append(lvl_text)
        
        # Paragraph properties for this level
        pPr = OxmlElement('w:pPr')
        
        # Indentation
        ind = OxmlElement('w:ind')
        ind.set(qn('w:left'), str(720 * level_idx))  # 0.5 inch per level
        ind.set(qn('w:hanging'), '360')  # 0.25 inch hanging
        pPr.append(ind)
        
        lvl.append(pPr)
        
        # Run properties for this level
        rPr = OxmlElement('w:rPr')
        fonts = OxmlElement('w:rFonts')
        fonts.set(qn('w:ascii'), FONT_NAME)
        fonts.set(qn('w:hAnsi'), FONT_NAME)
        fonts.set(qn('w:eastAsia'), FONT_NAME)
        rPr.append(fonts)
        lvl.append(rPr)
        
        abstract_num.append(lvl)
    
    numbering_xml.append(abstract_num)
    
    # Find existing num IDs to avoid conflicts
    existing_nums = numbering_xml.findall('.//w:num', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
    num_id = len(existing_nums) + 1
    
    # Create concrete numbering instance
    num = OxmlElement('w:num')
    num.set(qn('w:numId'), str(num_id))
    
    abstract_num_id_elem = OxmlElement('w:abstractNumId')
    abstract_num_id_elem.set(qn('w:val'), str(abstract_num_id))
    num.append(abstract_num_id_elem)
    
    numbering_xml.append(num)
    
    # Save changes
    numbering_part._numbering = numbering_xml
    
    return num_id


def create_heading_style(doc: Document, level: int, size: int, color: str, bold: bool, num_id: int):
    """Create or modify a heading style with numbering."""
    style_name = f'Heading {level}'
    
    if style_name not in doc.styles:
        doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
    
    style = doc.styles[style_name]
    style.font.name = FONT_NAME
    style.font.size = Pt(size)
    style.font.color.rgb = hex_to_rgb(color)
    style.font.bold = bold
    
    # Set paragraph spacing
    paragraph_format = style.paragraph_format
    paragraph_format.space_before = Pt(HEADING_SPACE_BEFORE)
    paragraph_format.space_after = Pt(HEADING_SPACE_AFTER)
    paragraph_format.keep_with_next = level < 3


def add_heading(doc: Document, text: str, level: int, num_id: int):
    """Add a heading to the document with numbering."""
    heading = doc.add_heading(text, level=level)
    
    # Apply specific styling based on level
    sizes = {1: HEADING_1_SIZE, 2: HEADING_2_SIZE, 3: HEADING_3_SIZE, 4: HEADING_4_SIZE}
    colors = {1: HEADING_1_COLOR, 2: HEADING_2_COLOR, 3: HEADING_3_COLOR, 4: HEADING_4_COLOR}
    bolds = {1: HEADING_1_BOLD, 2: HEADING_2_BOLD, 3: HEADING_3_BOLD, 4: HEADING_4_BOLD}
    
    size = sizes.get(level, HEADING_3_SIZE)
    color = colors.get(level, HEADING_3_COLOR)
    bold = bolds.get(level, HEADING_3_BOLD)
    
    set_paragraph_font(heading, font_size=size, font_color=color, bold=bold)
    
    # Add numbering - map heading level (1-4) to numbering level (0-3)
    numbering_level = level - 1
    add_numbering_to_paragraph(heading, num_id=num_id, level=numbering_level)


def add_code_block(doc: Document, code: str, language: Optional[str] = None):
    """Add a code block to the document with styling."""
    paragraph = doc.add_paragraph()
    
    if language:
        lang_run = paragraph.add_run(f"[{language}]\n")
        lang_run.font.name = FONT_NAME
        lang_run.font.size = Pt(8)
        lang_run.font.color.rgb = hex_to_rgb("#666666")
        lang_run.font.italic = True
    
    code_run = paragraph.add_run(code)
    code_run.font.name = CODE_FONT_NAME
    code_run.font.size = Pt(CODE_FONT_SIZE)
    code_run.font.color.rgb = hex_to_rgb("#000000")
    
    # Set paragraph shading
    pPr = paragraph._element.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), CODE_BACKGROUND_COLOR.replace('#', ''))
    pPr.append(shd)
    
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.left_indent = Inches(0.25)


def add_table(doc: Document, html_table):
    """Add an HTML table to the document with styling."""
    from bs4 import BeautifulSoup
    
    rows = []
    for tr in html_table.find_all('tr'):
        row_data = []
        for cell in tr.find_all(['td', 'th']):
            cell_text = ' '.join(cell.get_text().split())
            row_data.append(cell_text)
        if row_data:
            rows.append(row_data)
    
    if not rows:
        return
    
    num_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Table Grid'
    
    for row_idx, row_data in enumerate(rows):
        row = table.rows[row_idx]
        for col_idx, cell_text in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.text = cell_text
            
            if row_idx == 0:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = FONT_NAME
                        run.font.size = Pt(FONT_SIZE)
                        run.font.bold = True
                        run.font.color.rgb = hex_to_rgb(TABLE_HEADER_FONT_COLOR)
                
                tc = cell._element
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:fill'), TABLE_HEADER_BG_COLOR.replace('#', ''))
                tcPr.append(shd)
            else:
                for paragraph in cell.paragraphs:
                    set_paragraph_font(paragraph)
                    paragraph.paragraph_format.space_before = Pt(3)
                    paragraph.paragraph_format.space_after = Pt(3)
            
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            mar = OxmlElement('w:tcMar')
            for side in ['top', 'left', 'bottom', 'right']:
                mar_elem = OxmlElement(f'w:{side}')
                mar_elem.set(qn('w:w'), str(TABLE_CELL_PADDING))
                mar_elem.set(qn('w:type'), 'dxa')
                mar.append(mar_elem)
            tcPr.append(mar)


def process_html_element(doc: Document, element, num_id: int, in_code_block=False):
    """Process an HTML element and add it to the document."""
    from bs4 import BeautifulSoup, NavigableString
    
    if isinstance(element, NavigableString):
        text = str(element).strip()
        if text and not in_code_block:
            doc.add_paragraph(text)
        return
    
    tag_name = element.name if hasattr(element, 'name') else None
    
    if tag_name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
        level = int(tag_name[1])
        level = min(level, 4)
        text = element.get_text().strip()
        if text:
            add_heading(doc, text, level, num_id)
    
    elif tag_name == 'p':
        if element.find_parent('pre'):
            return
        
        text = element.get_text().strip()
        if text:
            paragraph = doc.add_paragraph(text)
            set_paragraph_font(paragraph)
    
    elif tag_name == 'pre':
        code_element = element.find('code')
        if code_element:
            code_text = code_element.get_text()
            language = None
            classes = code_element.get('class', [])
            for cls in classes:
                if cls.startswith('language-'):
                    language = cls.replace('language-', '')
                    break
            add_code_block(doc, code_text, language)
    
    elif tag_name == 'code':
        if not element.find_parent('pre'):
            code_text = element.get_text()
            paragraph = doc.add_paragraph()
            code_run = paragraph.add_run(code_text)
            code_run.font.name = CODE_FONT_NAME
            code_run.font.size = Pt(CODE_FONT_SIZE)
    
    elif tag_name == 'table':
        add_table(doc, element)
    
    elif tag_name in ['ul', 'ol']:
        list_type = 'bullet' if tag_name == 'ul' else 'number'
        for li in element.find_all('li', recursive=False):
            text = li.get_text().strip()
            if text:
                style = 'List Bullet' if list_type == 'bullet' else 'List Number'
                paragraph = doc.add_paragraph(text, style=style)
                set_paragraph_font(paragraph)
    
    elif tag_name == 'li':
        pass
    
    elif tag_name == 'hr':
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(12)
        paragraph.paragraph_format.space_after = Pt(12)
        pPr = paragraph._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '000000')
        pBdr.append(bottom)
        pPr.append(pBdr)
    
    elif tag_name == 'br':
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
    
    elif tag_name in ['strong', 'b']:
        text = element.get_text().strip()
        if text:
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(text)
            run.font.bold = True
            set_paragraph_font(paragraph)
    
    elif tag_name in ['em', 'i']:
        text = element.get_text().strip()
        if text:
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(text)
            run.font.italic = True
            set_paragraph_font(paragraph)
    
    elif tag_name == 'a':
        text = element.get_text().strip()
        href = element.get('href', '')
        if text:
            paragraph = doc.add_paragraph()
            if href:
                run = paragraph.add_run(f"{text} ({href})")
            else:
                run = paragraph.add_run(text)
            set_paragraph_font(paragraph)
    
    elif tag_name == 'blockquote':
        text = element.get_text().strip()
        if text:
            paragraph = doc.add_paragraph(text)
            set_paragraph_font(paragraph)
            pPr = paragraph._element.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            left = OxmlElement('w:left')
            left.set(qn('w:val'), 'single')
            left.set(qn('w:sz'), '12')
            left.set(qn('w:space'), '1')
            left.set(qn('w:color'), '7F7F7F')
            pBdr.append(left)
            pPr.append(pBdr)
            paragraph.paragraph_format.left_indent = Inches(0.5)
    
    else:
        for child in element.children:
            process_html_element(doc, child, num_id, in_code_block)


def convert_md_to_docx(input_path: str, output_path: str):
    """Convert a Markdown file to DOCX format."""
    input_file = Path(input_path)
    if not input_file.exists():
        raise FileNotFoundError(f"Input file not found: {input_path}")
    
    md_content = input_file.read_text(encoding='utf-8')
    html_content = parse_markdown(md_content)
    
    doc = Document()
    
    # Configure page setup
    section = doc.sections[0]
    section.page_width = Inches(PAGE_WIDTH)
    section.page_height = Inches(PAGE_HEIGHT)
    section.top_margin = Inches(PAGE_MARGIN_TOP)
    section.bottom_margin = Inches(PAGE_MARGIN_BOTTOM)
    section.left_margin = Inches(PAGE_MARGIN_LEFT)
    section.right_margin = Inches(PAGE_MARGIN_RIGHT)
    
    # Create multilevel numbering for headings
    num_id = create_multilevel_numbering(doc)
    
    # Create heading styles with numbering
    create_heading_style(doc, 1, HEADING_1_SIZE, HEADING_1_COLOR, HEADING_1_BOLD, num_id)
    create_heading_style(doc, 2, HEADING_2_SIZE, HEADING_2_COLOR, HEADING_2_BOLD, num_id)
    create_heading_style(doc, 3, HEADING_3_SIZE, HEADING_3_COLOR, HEADING_3_BOLD, num_id)
    create_heading_style(doc, 4, HEADING_4_SIZE, HEADING_4_COLOR, HEADING_4_BOLD, num_id)
    
    from bs4 import BeautifulSoup
    soup = BeautifulSoup(html_content, 'html.parser')
    
    for element in soup.children:
        process_html_element(doc, element, num_id)
    
    output_file = Path(output_path)
    output_file.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_file)
    
    print(f"Successfully converted '{input_path}' to '{output_path}'")


def main():
    """Main entry point."""
    if len(sys.argv) != 3:
        print("Usage: python md_to_docx.py <input.md> <output.docx>")
        print("\nExample:")
        print("  uv run --with python-docx --with markdown --with beautifulsoup4 scripts/md_to_docx.py docs/report_architecture.md docs/report_architecture.docx")
        sys.exit(1)
    
    input_path = sys.argv[1]
    output_path = sys.argv[2]
    
    try:
        convert_md_to_docx(input_path, output_path)
    except FileNotFoundError as e:
        print(f"Error: {e}")
        sys.exit(1)
    except Exception as e:
        print(f"Error converting file: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == "__main__":
    main()
