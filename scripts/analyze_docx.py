#!/usr/bin/env python3
"""
DOCX Analyzer - Inspects the structure and formatting of a Word document.
"""

import sys
from pathlib import Path
from docx import Document
from docx.enum.style import WD_STYLE_TYPE


def analyze_paragraph(paragraph, idx):
    """Analyze a paragraph's formatting."""
    info = {
        'index': idx,
        'text': paragraph.text[:100] + '...' if len(paragraph.text) > 100 else paragraph.text,
        'style': paragraph.style.name if paragraph.style else 'None',
        'is_heading': False,
        'heading_level': None,
        'font_name': None,
        'font_size': None,
        'font_color': None,
        'bold': False,
        'italic': False,
        'has_numbering': False,
        'alignment': paragraph.alignment.name if paragraph.alignment else 'None',
    }
    
    # Check if it's a heading style
    if paragraph.style.name.startswith('Heading'):
        info['is_heading'] = True
        info['heading_level'] = int(paragraph.style.name.split()[-1])
    
    # Check for numbering
    pPr = paragraph._element.pPr
    if pPr is not None:
        numPr = pPr.numPr
        if numPr is not None:
            info['has_numbering'] = True
    
    # Get font info from first run
    if paragraph.runs:
        run = paragraph.runs[0]
        info['font_name'] = run.font.name
        info['font_size'] = run.font.size.pt if run.font.size else None
        if run.font.color and run.font.color.rgb:
            info['font_color'] = run.font.color.rgb
        info['bold'] = run.font.bold
        info['italic'] = run.font.italic
    
    return info


def analyze_document(doc_path: str):
    """Analyze a Word document and print its structure."""
    doc = Document(doc_path)
    
    print("=" * 80)
    print(f"DOCUMENT ANALYSIS: {doc_path}")
    print("=" * 80)
    
    # Document statistics
    print(f"\n📊 DOCUMENT STATISTICS")
    print(f"   Total paragraphs: {len(doc.paragraphs)}")
    print(f"   Total sections: {len(doc.sections)}")
    print(f"   Available styles: {len(doc.styles)}")
    
    # Page setup
    print(f"\n📄 PAGE SETUP")
    section = doc.sections[0]
    print(f"   Page width: {section.page_width.inches} inches")
    print(f"   Page height: {section.page_height.inches} inches")
    print(f"   Top margin: {section.top_margin.inches} inches")
    print(f"   Bottom margin: {section.bottom_margin.inches} inches")
    print(f"   Left margin: {section.left_margin.inches} inches")
    print(f"   Right margin: {section.right_margin.inches} inches")
    
    # Analyze all paragraphs
    print(f"\n📝 PARAGRAPH ANALYSIS")
    print("-" * 80)
    
    headings = []
    paragraphs_with_issues = []
    font_variations = {}
    size_variations = {}
    
    for idx, paragraph in enumerate(doc.paragraphs):
        info = analyze_paragraph(paragraph, idx)
        
        # Track headings
        if info['is_heading']:
            headings.append(info)
        
        # Track font variations
        font_key = f"{info['font_name']}_{info['font_size']}_{info['font_color']}"
        font_variations[font_key] = font_variations.get(font_key, 0) + 1
        
        if info['font_size']:
            size_key = str(info['font_size'])
            size_variations[size_key] = size_variations.get(size_key, 0) + 1
        
        # Identify issues
        issues = []
        if info['is_heading']:
            if info['font_name'] != 'Arial':
                issues.append(f"Font is '{info['font_name']}' not 'Arial'")
            if info['font_size'] not in [16, 14, 13, 12]:
                issues.append(f"Size {info['font_size']}pt may not match standard")
            if info['font_color'] and str(info['font_color']) != '000000':
                issues.append(f"Color is not black ({info['font_color']})")
        else:
            if info['font_name'] and info['font_name'] != 'Arial':
                issues.append(f"Font is '{info['font_name']}' not 'Arial'")
            if info['font_size'] and info['font_size'] != 11:
                issues.append(f"Size is {info['font_size']}pt not 11pt")
            if info['font_color'] and str(info['font_color']) != '000000':
                issues.append(f"Color is not black ({info['font_color']})")
        
        if issues:
            paragraphs_with_issues.append({**info, 'issues': issues})
        
        # Print heading structure
        if info['is_heading']:
            indent = "  " * (info['heading_level'] - 1)
            print(f"   {indent}H{info['heading_level']}: {info['text'][:60]}")
            if info['has_numbering']:
                print(f"      ✓ Has numbering")
            else:
                print(f"      ✗ Missing numbering")
            print(f"      Font: {info['font_name']}, Size: {info['font_size']}pt, Color: {info['font_color']}, Bold: {info['bold']}")
    
    # Summary of font variations
    print(f"\n🎨 FONT VARIATIONS FOUND")
    print("-" * 80)
    for font_key, count in sorted(font_variations.items(), key=lambda x: -x[1])[:10]:
        print(f"   {count}x - {font_key}")
    
    print(f"\n📏 FONT SIZE DISTRIBUTION")
    print("-" * 80)
    for size, count in sorted(size_variations.items(), key=lambda x: -x[1]):
        print(f"   {size}pt: {count} paragraphs")
    
    # Issues summary
    print(f"\n⚠️  FORMATTING ISSUES FOUND")
    print("-" * 80)
    print(f"   Total paragraphs with issues: {len(paragraphs_with_issues)}")
    
    if paragraphs_with_issues:
        print(f"\n   First 20 issues:")
        for info in paragraphs_with_issues[:20]:
            print(f"   [{info['index']}] {info['text'][:50]}...")
            for issue in info['issues']:
                print(f"      - {issue}")
    
    # Heading structure
    print(f"\n📋 HEADING STRUCTURE")
    print("-" * 80)
    for info in headings:
        indent = "  " * (info['heading_level'] - 1)
        numbering_status = "✓" if info['has_numbering'] else "✗"
        print(f"   {indent}{numbering_status} H{info['heading_level']}: {info['text'][:60]}")
    
    return {
        'total_paragraphs': len(doc.paragraphs),
        'headings': headings,
        'issues': paragraphs_with_issues,
        'font_variations': font_variations,
        'size_variations': size_variations,
    }


if __name__ == "__main__":
    if len(sys.argv) != 2:
        print("Usage: python analyze_docx.py <document.docx>")
        sys.exit(1)
    
    doc_path = sys.argv[1]
    analyze_document(doc_path)
