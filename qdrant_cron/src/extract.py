import io
import os
import magic
import fitz
from docx import Document
from pathlib import Path
from typing import Optional
from util.logger import get_logger

logger = get_logger()


def _extract_pdf_text(doc: fitz.Document):
    text = ""
    for page in doc:
        page_text = page.get_text("text")
        assert isinstance(page_text, str)
        text += page_text
    return text


def extract_text(*, file_path: Optional[str | Path] = None, file_bytes: Optional[bytes] = None):
    """
    Detect file type by magic bytes.
    Either file_path or file_bytes must be provided.

    When providing file_bytes, it is recommended to just read first 2048 bytes of the file where possible.
    """
    mime = magic.Magic(mime=True)
    blob = bytes()

    if file_bytes:
        blob = file_bytes
    else:
        assert file_path
        with open(file_path, "rb") as f:
            blob = f.read(2048)

    file_type = mime.from_buffer(blob)
    print(f"Detected: {file_type} {f'for {os.path.basename(file_path)}' if file_path else ''}")

    try:
        # PDF Parsing
        if file_type == "application/pdf":
            text = ""
            if file_bytes:
                with fitz.open(stream=file_bytes, filetype="pdf") as doc:
                    return _extract_pdf_text(doc)
            else:
                assert file_path
                with fitz.open(file_path) as doc:
                    return _extract_pdf_text(doc)

            return text

        # DOCX Parsing
        elif (
            file_type
            == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ):
            if file_bytes:
                doc = Document(io.BytesIO(file_bytes))
            else:
                assert file_path
                doc = Document(str(file_path))

            # Joins all paragraphs with newlines
            return "\n".join([para.text for para in doc.paragraphs])

        # Plain Text / Fallback
        elif "text/" in file_type or file_type == "application/x-empty":
            if file_bytes:
                return file_bytes.decode("utf-8")
            else:
                assert file_path
                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    return f.read()

        else:
            return f"Unsupported file type: {file_type}"

    except Exception as e:
        return f"Error processing {file_path}: {str(e)}"
